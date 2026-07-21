#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 EnzyMiner Pro 重要源代码 Word 文档。

运行：
    python scripts/build_source_code_doc.py

输出：
    docs/EnzyMiner_Pro_重要源代码文档_V1.1.1.docx

文档包含项目概览、技术架构、源文件清单，以及项目中重要代码/配置文件的完整内容。
默认排除 node_modules、dist、日志、运行时任务数据和大体积示例数据，避免把构建产物或第三方依赖误当作项目源代码。
"""

from __future__ import annotations

import argparse
import hashlib
import os
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUTPUT = ROOT / "docs" / "EnzyMiner_Pro_重要源代码文档_V1.1.1.docx"

# 项目代码和可执行配置。按目录明确纳入，保证文档内容稳定、可复现。
SOURCE_GLOBS = (
    "backend/**/*.py",
    "backend/**/*.mjs",
    "scripts/**/*.py",
    "scripts/**/*.sh",
    "src/**/*.tsx",
    "src/**/*.ts",
    "src/**/*.css",
    "test/**/*.mjs",
)
ROOT_FILES = (
    ".env.example",
    ".gitignore",
    ".npmrc",
    "index.html",
    "package.json",
    "package-lock.json",
    "requirements.txt",
    "start.sh",
    "stop.sh",
    "tsconfig.json",
    "vite.config.ts",
)

LANGUAGE_BY_SUFFIX = {
    ".css": "CSS",
    ".html": "HTML",
    ".js": "JavaScript",
    ".json": "JSON",
    ".mjs": "JavaScript (ESM)",
    ".py": "Python",
    ".sh": "Shell",
    ".ts": "TypeScript",
    ".tsx": "TypeScript + React JSX",
    ".txt": "Text",
}

FONT_EA = "Microsoft YaHei"
FONT_CODE = "Consolas"
BLUE = (31, 56, 100)
GRAY = (90, 90, 90)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, size: float = 9.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    run.font.name = FONT_CODE if not bold else FONT_EA
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_EA)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, *, font: str = FONT_EA, size: float = 10.5, bold: bool = False,
                 color: tuple[int, int, int] | None = None) -> None:
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_EA)


def add_paragraph(doc: Document, text: str = "", *, size: float = 10.5,
                  bold: bool = False, align=None, color=None, after: float = 6):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    if align is not None:
        paragraph.alignment = align
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_heading(doc: Document, text: str, level: int):
    paragraph = doc.add_heading(level=level)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, size={1: 18, 2: 14, 3: 11.5}.get(level, 10.5), bold=True, color=BLUE)
    return paragraph


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=GRAY)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run2 = paragraph.add_run(" 页")
    set_run_font(run2, size=9, color=GRAY)


def add_toc(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = r'TOC \o "1-3" \h \z \u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "打开文档后在此处右键并选择“更新域”生成目录。"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(separate)
    run._r.append(placeholder)
    run._r.append(end)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_EA
    normal.font.size = Pt(10.5)
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), FONT_EA)

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.text = ""
        add_page_number(footer)


def collect_source_files() -> list[Path]:
    files: set[Path] = set()
    for pattern in SOURCE_GLOBS:
        files.update(path for path in ROOT.glob(pattern) if path.is_file())
    for relative in ROOT_FILES:
        path = ROOT / relative
        if path.is_file():
            files.add(path)
    return sorted(files, key=lambda path: str(path.relative_to(ROOT)).lower())


def read_source(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def file_info(path: Path) -> dict[str, object]:
    content = read_source(path)
    raw = path.read_bytes()
    return {
        "path": path.relative_to(ROOT).as_posix(),
        "language": LANGUAGE_BY_SUFFIX.get(path.suffix.lower(), "Source"),
        "lines": len(content.splitlines()),
        "bytes": len(raw),
        "sha256": hashlib.sha256(raw).hexdigest()[:16],
        "content": content,
    }


def add_table(doc: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[index], header, bold=True)
        set_cell_shading(table.rows[0].cells[index], "D9E2F3")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            set_cell_text(cells[index], str(value), size=9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc: Document, content: str) -> None:
    # 一个源代码行对应一个 Word 段落，保留空行、缩进和行号，便于检索与评审。
    lines = content.splitlines()
    if not lines:
        lines = [""]
    for number, line in enumerate(lines, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.left_indent = Cm(0.1)
        paragraph.paragraph_format.keep_together = True
        run = paragraph.add_run(f"{number:>5}  {line}")
        set_run_font(run, font=FONT_CODE, size=7.3)


def build_document(output: Path) -> tuple[Path, int]:
    output.parent.mkdir(parents=True, exist_ok=True)
    source_files = [file_info(path) for path in collect_source_files()]
    doc = Document()
    configure_document(doc)

    # 封面
    add_paragraph(doc, "EnzyMiner Pro", size=26, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=BLUE, after=12)
    add_paragraph(doc, "重要源代码文档", size=20, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, color=BLUE, after=20)
    add_paragraph(doc, "高通量酶挖掘全栈工具平台", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, after=8)
    add_paragraph(doc, f"生成日期：{date.today().isoformat()}", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, after=4)
    add_paragraph(doc, f"纳入文件：{len(source_files)} 个", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, after=4)
    add_paragraph(doc, "本文件由 scripts/build_source_code_doc.py 自动生成。", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=GRAY, after=20)
    doc.add_page_break()

    # 文档说明
    add_heading(doc, "1. 文档说明", 1)
    add_paragraph(doc, "本文件用于集中记录 EnzyMiner Pro 项目中承担核心业务、计算流程、前端交互、后端接口、自动化脚本和构建配置职责的源代码。除项目概览和文件索引外，后续章节按文件逐行收录完整内容，可用于代码审阅、版本归档、离线查阅和软件交付。")
    add_paragraph(doc, "收录范围：backend、scripts、src、test 目录中的代码文件，以及项目根目录中的启动脚本、构建配置、依赖清单和环境变量示例。排除 node_modules、dist、日志、运行时任务数据及示例结果数据；这些内容属于第三方依赖、构建产物或运行数据，不作为项目核心源代码收录。")

    add_heading(doc, "2. 项目概览与技术架构", 1)
    add_paragraph(doc, "EnzyMiner Pro 是面向高通量酶挖掘的全栈生物信息学平台，集成 BLAST、HMMER 和 Compare 模块，覆盖参考序列管理、同源搜索、多序列比对、位点打分、聚类去冗余、相似性计算、网络分析、性质预测、候选推荐和任务报告生成等流程。")
    add_table(doc, ["层次", "主要技术/文件", "职责"], [
        ["前端", "React + TypeScript + Vite；src/", "流水线页面、任务状态、序列比对、网络可视化、性质预测与推荐交互"],
        ["后端", "Express；backend/server.mjs、taskReport.mjs", "REST API、任务管理、文件读写、计算调度、结果聚合和报告导出"],
        ["计算层", "Python；pipeline.py、biopython_pairwise_similarity.py", "序列获取、BLAST/HMMER 流程、比对、评分、聚类及相似性计算"],
        ["辅助脚本", "scripts/", "NCBI/UniProt 注释、服务启动停止和冒烟测试"],
        ["测试与构建", "test/、package.json、vite.config.ts", "集成测试、依赖管理、开发/生产构建配置"],
    ])

    add_heading(doc, "3. 源文件索引", 1)
    rows = []
    for info in source_files:
        rows.append([info["path"], info["language"], info["lines"], info["bytes"], info["sha256"]])
    add_table(doc, ["文件", "语言", "行数", "字节数", "SHA-256 前16位"], rows)
    doc.add_page_break()

    # 完整源代码
    add_heading(doc, "4. 重要源代码全文", 1)
    add_paragraph(doc, "以下内容为生成文档时从工作区读取的源文件快照。每个文件单独成节，代码左侧为原始行号；文件路径、语言、大小和 SHA-256 前缀用于定位和校验。")
    for index, info in enumerate(source_files):
        add_heading(doc, f"4.{index + 1} {info['path']}", 2)
        add_table(doc, ["属性", "值"], [
            ["语言", info["language"]],
            ["行数", info["lines"]],
            ["大小", f"{info['bytes']} bytes"],
            ["SHA-256 前16位", info["sha256"]],
        ])
        add_code_block(doc, info["content"])
        if index != len(source_files) - 1:
            doc.add_page_break()

    add_heading(doc, "5. 文档生成与更新", 1)
    add_paragraph(doc, "在项目根目录执行以下命令即可重新生成本文件；脚本会以当前工作区内容为准，自动刷新文件清单、行数、校验值和源代码全文：")
    command = doc.add_paragraph()
    command.paragraph_format.left_indent = Cm(0.5)
    command.paragraph_format.space_after = Pt(8)
    run = command.add_run("python scripts/build_source_code_doc.py")
    set_run_font(run, font=FONT_CODE, size=9)
    add_paragraph(doc, "也可以通过 --output 指定其他输出路径，例如：python scripts/build_source_code_doc.py --output /tmp/enzyminer-source.docx。")

    # 让 Word 打开时提示更新目录/域。
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)
    doc.save(output)
    return output, len(source_files)


def main() -> None:
    parser = argparse.ArgumentParser(description="生成 EnzyMiner Pro 重要源代码 Word 文档")
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT, help="输出 DOCX 路径")
    args = parser.parse_args()
    output = args.output if args.output.is_absolute() else ROOT / args.output
    path, count = build_document(output)
    print(f"已生成：{path}")
    print(f"收录源文件：{count} 个")


if __name__ == "__main__":
    main()
